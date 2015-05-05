#coding=utf-8
from django.http import HttpResponseRedirect, HttpResponse
from django.shortcuts import render,get_object_or_404
from reports.models import *
from django.http import Http404
from django.core.urlresolvers import reverse

import pdb
import datetime,time,os,sys
from django.core.files import File
from generic import *
#import pythonwhois
import urllib
import json
from django.conf import settings

import base64


mrtg_image_dic={}
dnsla_image_dic={}
show_persion_nb=""
show_domain_nb=""
show_ip_nb=""
show_process_nb=""
attact_nodes={}

# Create your views here.
def index(request):
    latest_report_list = Report.objects.order_by('-create_time')[:10]
    context = {'latest_report_list': latest_report_list}
    return render(request, 'reports/index.html', context)


def create(request):

    # if this is a POST request we need to process the form data
    
    global mrtg_image_dic
    global dnsla_image_dic
    if request.method == 'POST':
        #从request的数据中,初始化一个表单
        form = ReportForm(request.POST)
        #表单数据验证
        if form.is_valid():
            #验证成功,存入数据库
            ret=save_form(form,0)
            if ret['success']:
                #return HttpResponseRedirect('/reports/')
                #重定向到预览页面
                return view(request,ret['report_id'])
            else:
                #保存失败 返回错误
                ret['is_form']=False
                return render(request,'reports/create.html',ret)
        else:
            try:
                nodeidlist=form.cleaned_data['mrtg_image']
                if '0' in nodeidlist:
                    nodeidlist.remove('0')
            except:
                nodeidlist=[]
                
            nodes=Node.objects.all()
            for i in xrange(0,len(nodes)):
                dnsla_image_dic[str(nodes[i].nb)]={'name':nodes[i].name,'dname':{'x':[],'y':[]},'addr':{'x':[],'y':[]}}
            return render(request, 'reports/create.html', {
                'form': form,
                'mrtg_image_dic':mrtg_image_dic,
                'dnsla_image_dic':dnsla_image_dic,
                'is_form':True,
                'is_create':True,
                'is_valid':False,
                'nodeidlist':nodeidlist,
            })
    # if a GET (or any other method) we'll create a blank form
    else:
        #生成一个空的表单,非绑定
        form = ReportForm()
        #生成全局图像字典变量
        mrtg_image_dic=get_mrtg_image()
        #dnsla_image_dic=get_dnsla_image()
        #dnsla_image_dic={}
        nodes=Node.objects.all()
        for i in xrange(0,len(nodes)):
            dnsla_image_dic[str(nodes[i].nb)]={'name':nodes[i].name,'dname':{'x':[],'y':[]},'addr':{'x':[],'y':[]}}

    return render(request, 'reports/create.html', {
            'form': form,
            'mrtg_image_dic':mrtg_image_dic,
            'dnsla_image_dic':dnsla_image_dic,
            'is_form':True,
            'is_create':True,
            'is_valid':True,
            })


        
def change(request,report_id):

    global show_persion_nb
    global show_domain_nb
    global show_ip_nb
    global show_process_nb
    global dnsla_image_dic
    global mrtg_image_dic
    global attact_nodes
    
    if request.method == 'POST':
        form = ReportForm(request.POST)
        
        if form.is_valid():
            ret=save_form(form,report_id)
            
            if ret['success']:
                #return HttpResponseRedirect('/reports/')
                #重定向到预览页面
                return view(request,report_id)
            else:
                #保存失败 返回错误
                ret['is_form']=False
                return render(request,'reports/create.html',ret)
        #change验证未通过
        else:
            try:
                nodeidlist=form.cleaned_data['mrtg_image']
                if '0' in nodeidlist:
                    nodeidlist.remove('0')
            except:
                nodeidlist=[]
                
            return render(request, 'reports/create.html', {
            'form': form,
            'is_form':True,
            'is_create':False,
            'report_id':report_id,
            "attact_nodes":attact_nodes,
            "show_persion_nb":show_persion_nb,
            "show_domain_nb":show_domain_nb,
            "show_ip_nb":show_ip_nb,
            "show_process_nb":show_process_nb,
            "dnsla_image_dic":dnsla_image_dic,
            "mrtg_image_dic":mrtg_image_dic,
            "is_change_valid":False,
            "nodeidlist":nodeidlist,
            })
    else:
        report = get_object_or_404(Report, pk=report_id)
        attact_nodes=report.attact_nodes.all()
        
        mrtg_image_index=[]
        for an in attact_nodes:
            try:
                node=Node.objects.get(nb=an.nb)
                an.name=node.name
                an.abbr=node.abbr
                mrtg_image_index.append(str(an.nb))
                dnsla_image_dic[str(node.nb)]={'name':node.name,'dname':{'x':[],'y':[]},'addr':{'x':[],'y':[]}}
            except Exception,e:
                pass

        domain_str=":::".join(domain.domain_name for domain in report.domains.all())
        ip_str=":::".join(ip.addr for ip in report.ips.all())
        change_dic={
        "title":report.title,
        "start_time":report.start_time,
        "end_time":report.end_time,
        "persion":report.persion,
        "domain":domain_str,
        "ip":ip_str,
        "abstract":report.abstract,
        "process":report.process,
        "mrtg_image":mrtg_image_index,
        }

        form = ReportForm(initial=change_dic)

        #刷新全局变量
        mrtg_image_dic=get_mrtg_image()
        nodes=Node.objects.all()
        for i in xrange(0,len(nodes)):
            dnsla_image_dic[str(nodes[i].nb)]={'name':nodes[i].name,'dname':{'x':[],'y':[]},'addr':{'x':[],'y':[]}}
        show_persion_nb=get_nb_of_str(report.persion,":::")
        show_domain_nb=get_nb_of_str(domain_str,":::")
        show_ip_nb=get_nb_of_str(ip_str,":::")
        show_process_nb=get_nb_of_str(report.process,":::")

    return render(request, 'reports/create.html', {
        'form': form,
        'is_form':True,
        'is_create':False,
        'report_id':report_id,
        "attact_nodes":attact_nodes,
        "show_persion_nb":show_persion_nb,
        "show_domain_nb":show_domain_nb,
        "show_ip_nb":show_ip_nb,
        "show_process_nb":show_process_nb,
        "dnsla_image_dic":dnsla_image_dic,
        "mrtg_image_dic":mrtg_image_dic,
        "is_change_valid":True,
        })

def get_nb_of_str(string,sp):
    lis=string.split(sp)
    ret=0
    for i in xrange(0,len(lis)):
        if lis[i]!="":
            ret=i+1
    return ret

def save_form(form,report_id):
    global mrtg_image_dic
    global dnsla_image_dic
    ret={'success':True,'error_message':'','report_id':report_id}
    #新建一个报告
    if report_id==0:
        try:
            start_time=form.cleaned_data['start_time']
            end_time=form.cleaned_data['end_time']

            report=Report(
                title=form.cleaned_data['title'],
                create_time=datetime.datetime.now(),
                persion=form.cleaned_data['persion'],
                abstract=form.cleaned_data['abstract'],
                process=form.cleaned_data['process'],
                start_time=start_time,
                end_time=end_time
            )
            report.save()
        except Exception,e:
            ret['success']=False
            ret['error_message']=str(e)
            return ret
    #修改已有的报告
    else:
        try:
            report=get_object_or_404(Report, pk=report_id)
            #这两个变量后面要用到
            start_time=form.cleaned_data['start_time']
            end_time=form.cleaned_data['end_time']
            
            report.start_time=start_time
            report.end_time=end_time
            report.title=form.cleaned_data['title']
            #create_time=datetime.datetime.now()
            report.persion=form.cleaned_data['persion']
            report.abstract=form.cleaned_data['abstract']
            report.process=form.cleaned_data['process']
        
            report.save()
        except Exception,e:
            ret['success']=False
            ret['error_message']=str(e)
            return ret
    
    #处理domain
    origin_ele=set(domain.domain_name for domain in report.domains.all())
    new_ele=set(form.cleaned_data['domain'].split(':::'))
    
    del_ele=list(origin_ele-new_ele)
    add_ele=list(new_ele-origin_ele)
    
    for domain in del_ele:
        dom=Domain.objects.get(domain_name=domain)
        dom.report.remove(report)
        if len(dom.report.all())==0:
            dom.delete()

    for domain in add_ele:
        if domain !='':
            try:
                dom=Domain.objects.get(domain_name=domain)
            except Domain.DoesNotExist:
                dom=Domain(domain_name=domain)
            """
            try:
                #whois 查询
                result=pythonwhois.get_whois(domain)
                dom.status=';'.join(result['status'])
                dom.registrant=result['contacts']['registrant']['name']
            except Exception,e:
            """
            dom.status=''
            dom.registrant=''
            dom.save()
            dom.report.add(report)

    #处理IP
    origin_ele=set(ip.addr for ip in report.ips.all())
    new_ele=set(form.cleaned_data['ip'].split(':::'))
    
    del_ele=list(origin_ele-new_ele)
    add_ele=list(new_ele-origin_ele)
    
    #处理要删除的元素
    for ip in del_ele:
        ip_obj=Ip.objects.get(addr=ip)
        ip_obj.report.remove(report)
        if len(ip_obj.report.all())==0:
            ip_obj.delete()
    
    #处理要添加的元素
    for ip in add_ele:
        if ip !='':
            try:
                ip_obj=Ip.objects.get(addr=ip)
                ip_obj.report.add(report)
            except Ip.DoesNotExist:
                ip_obj=Ip(addr=ip)
                ip_obj.save()
                ip_obj.report.add(report)

    
    #修改时删除所有攻击节点,重新赋值
    for an in report.attact_nodes.all():
        an.delete()
    for i in form.cleaned_data['mrtg_image']:
        if mrtg_image_dic.has_key(str(i)) and dnsla_image_dic.has_key(str(i)):  
            try:
                img_name=mrtg_image_dic[str(i)]
                #dnsla图像以json数据格式保存,前段echarts库图形显示
                dnsla_data=dnsla_image_dic[str(i)]
                #计算该节点平均qps和峰值qps
                #start_time<end_time<end_time
                node_qps=Qps.objects.filter(node_id=int(i),end_time__gt=start_time,end_time__lt=end_time)
                inqps_list=[]
                outqps_list=[]
                for nq in node_qps:
                    inqps_list.append(nq.in_qps)
                    outqps_list.append(nq.out_qps)
                
                max_qps=max_of_list(inqps_list+outqps_list)
                average_qps=average_of_list(inqps_list+outqps_list)
                att=Attact_node(
                    report=report,
                    nb=int(i),max_qps=max_qps,
                    average_qps=average_qps,
                    picture=img_name,
                    dnsla_json=json.dumps(dnsla_data)
                )
                att.save()
                
            except Exception,e:
                ret['success']=False
                ret['error_message']=str(e)
                return ret
    
    ret['report_id']=report.id
    return ret

def view(request,report_id):
    report = get_object_or_404(Report, pk=report_id)
    create_time=str(report.create_time.date())
    persions=report.persion.split(':::')
    
    year=report.start_time.year
    month=report.start_time.month
    day=report.start_time.day
    end_day=report.end_time.day
    
    start_time=str(report.start_time.time())
    end_time=str(report.end_time.time())
    
    #nodes=report.nodes.all()
    
    #pdb.set_trace()
    attact_nodes=report.attact_nodes.all()
    for an in attact_nodes:
        try:
            node=Node.objects.get(nb=an.nb)
            an.name=node.name
            an.abbr=node.abbr
        except:
            continue
    
    
    context={
        'report':report,
        'create_time':create_time,
        'persions':persions,
        'year':year,'month':month,'day':day,'end_day':end_day,
        'start_time':start_time,'end_time':end_time,
        'attact_nodes':attact_nodes,
        'processes':report.process.split(':::'),
        
    }
    return render(request, 'reports/view.html', context)

def detail(request):
    pass

def get_mrtg_image():
    #格式化时间
    end_time=int(time.time())
    start_time=end_time-86400
	
    now=time.localtime(time.time())
    comment_now=time.strftime('%Y-%m-%d %H\:%M\:%S',now)
    file_time=time.strftime('%Y%m%d%H%M%S',now)
	
    rrd_tool=settings.RRD_TOOL
    rrd_file_path=settings.RRD_FILE_PATH
    rrd_files=os.popen("find "+ rrd_file_path +" -name *.rrd").read().strip().split('\n')

    ret={}
    for rrd_file in rrd_files:
        
        image_file_name_key=rrd_file.replace(os.path.join(rrd_file_path,'cnqps_'),'').replace('.rrd','')
        try:
            #image_file_num=name_to_num[image_file_name_key]
            node=Node.objects.get(abbr=image_file_name_key)
            image_file_num=str(node.nb)
            node_name=node.name
        except:
            image_file_num='-1'
            node_name='空'
        image_file_name=image_file_name_key+'_'+image_file_num+'_'+file_time+'.png'
        image_file_path=os.path.join(mrtg_image_save_dir,image_file_name)
        #pdb.set_trace()
        os.popen('%s graph %s  --start %s --end %s         \
			     --title "Daily\' Graph (5 Minutes Average) "      \
			     --vertical-label "Bits Per second"               \
			     --height 400                      \
			     --width  1000                     \
			     --color "BACK#CCCCCC"             \
			     --color "CANVAS#CCFFFF"           \
			     --slope-mode                      \
			     --lower-limit 0                   \
			     DEF:value1=%s:ds0:AVERAGE   \
			     DEF:value2=%s:ds1:AVERAGE   \
			     DEF:value3=%s:ds0:MAX       \
			     DEF:value4=%s:ds1:MAX       \
			     COMMENT:"               Max              Average           Last\\n"  \
			     AREA:value1#00FF00:" In\\:"                \
			     GPRINT:value1:MAX:"%%10.2lf %%Sb/s"         \
			     GPRINT:value1:AVERAGE:"%%10.2lf %%sb/s"     \
			     GPRINT:value1:"LAST:%%10.2lf %%Sb/s"        \
			     COMMENT:" \\n"                             \
			     LINE3:value2#0000FF:"Out\\:"               \
			     GPRINT:value2:MAX:"%%10.2lf %%Sb/s"         \
			     GPRINT:value2:AVERAGE:"%%10.2lf %%Sb/s"     \
			     GPRINT:value2:LAST:"%%10.2lf %%Sb/S"        \
			     COMMENT:" \\n"                             \
			     COMMENT:"Last update\\: %s"'%(rrd_tool,image_file_path,start_time,end_time,rrd_file,rrd_file,rrd_file,rrd_file,str(comment_now)))
        
        ret[image_file_num]="/media/mrtg/"+image_file_name
    return ret

def dnsla_get_json_data(url,nodeid,topn,data_type):

    if settings.DEBUG:
        if data_type=='dname':
            json_data='{"nodeid":1,"num":30,"type":"domain","rank":[{"dname":"in-addr.arpa.","count":"643"},{"dname":"in-addr.arpa.","count":"350"},{"dname":"in-addr.arpa.","count":"110"},{"dname":"teny.cn.","count":"70"},{"dname":"teny.cn.","count":"51"},{"dname":"360.cn.","count":"44"},{"dname":"dynamic.163data.com.cn.","count":"40"},{"dname":"CNiNfO.Com.cN.","count":"37"},{"dname":"dns.cn.","count":"33"},{"dname":"online.sh.cn.","count":"33"},{"dname":"ns.ptt.js.cn.","count":"32"},{"dname":"360.cn.","count":"29"},{"dname":"sina.com.cn.","count":"26"},{"dname":"zol.com.cn.","count":"25"},{"dname":"ccgslb.com.cn.","count":"23"},{"dname":"ce.cn.","count":"23"},{"dname":"bta.net.cn.","count":"22"},{"dname":"ccgslb.com.cn.","count":"19"},{"dname":"online.sh.cn.","count":"17"},{"dname":"cninfo.com.cn.","count":"17"},{"dname":"gRandClouD.cn.","count":"16"},{"dname":"dns.com.cn.","count":"16"},{"dname":"dynamic.163data.com.cn.","count":"15"},{"dname":"amazon.cn.","count":"15"},{"dname":"sdjnptt.net.cn.","count":"15"},{"dname":"cnnic.cn.","count":"15"},{"dname":"uc.cn.","count":"14"},{"dname":"3g.cn.","count":"13"},{"dname":"yihaodian.com.cn.","count":"13"},{"dname":"mq.hl.cn.","count":"12"}]}'
        else:
            json_data='{"nodeid":1,"num":30,"type":"addr","rank":[{"addr":"101.226.160.0","count":"917"},{"addr":"220.181.108.0","count":"415"},{"addr":"174.36.22.3","count":"167"},{"addr":"180.153.229.0","count":"255"},{"addr":"195.2.240.2","count":"87"},{"addr":"180.97.35.0","count":"206"},{"addr":"74.125.46.0","count":"82"},{"addr":"42.156.207.0","count":"146"},{"addr":"173.194.98.0","count":"53"},{"addr":"61.139.113.0","count":"118"},{"addr":"104.156.251.22","count":"48"},{"addr":"115.239.212.0","count":"63"},{"addr":"220.181.108.0","count":"364"},{"addr":"212.54.41.0","count":"43"},{"addr":"202.101.173.0","count":"58"},{"addr":"123.125.71.0","count":"302"},{"addr":"91.210.72.56","count":"33"},{"addr":"202.96.209.0","count":"50"},{"addr":"82.96.64.2","count":"32"},{"addr":"111.206.36.0","count":"258"},{"addr":"101.226.3.85","count":"45"},{"addr":"81.175.171.0","count":"25"},{"addr":"180.153.229.0","count":"246"},{"addr":"222.217.39.0","count":"44"},{"addr":"195.177.123.1","count":"18"},{"addr":"211.98.71.0","count":"120"},{"addr":"180.137.252.0","count":"36"},{"addr":"220.181.12.0","count":"113"},{"addr":"195.94.224.4","count":"17"},{"addr":"61.188.16.0","count":"33"}]}'
    
    else:
        f = urllib.urlopen(url % str(nodeid))
        json_data=f.read()
    #去除无用字符
    #BOM_UTF8 '\xef\xbb\xbf' 
    data=json.loads(json_data)
    x_list=[]
    y_list=[]
    

    this_type=data['type']
    if data['rank']==None:
        return [],[]
        
    rank=sorted(data['rank'],key=lambda e:int(e['count']),reverse=True)
    
    #data['rank'].sort(key=lambda e:int(e['count']),reverse=True)
    for i in xrange(0,len(rank)):
        if i>topn:
            break
        this_rank=rank[i]
        y_list.append(this_rank[data_type])
        x_list.append(this_rank['count'])
    y_list.reverse()
    x_list.reverse()
    return x_list,y_list
    
def get_dnsla_image():
    ret_dic={}
    for i in xrange(0,32):
        ret_dic[str(i)]={}
        x_list,y_list=dnsla_get_json_data(url=dnsla_dname_url,nodeid=i,topn=10,data_type='dname')
        ret_dic[str(i)]['dname']={'x':x_list,'y':y_list}
        x_list,y_list=dnsla_get_json_data(url=dnsla_addr_url,nodeid=i,topn=10,data_type='addr')
        ret_dic[str(i)]['addr']={'x':x_list,'y':y_list}
    return ret_dic

def getDnslaCurrentData(request,nodeid):
    
    if nodeid=='0':
        return  HttpResponse(json.dumps(""))
    global dnsla_image_dic
    if request.method == 'GET':
        nodeid=int(nodeid)
        #获取dnsla数据
        x_dname,y_dname=dnsla_get_json_data(url=dnsla_dname_url,nodeid=nodeid,topn=10,data_type='dname')
        x_addr,y_addr=dnsla_get_json_data(url=dnsla_addr_url,nodeid=nodeid,topn=10,data_type='addr')
        #存储到全局变量中
        dnsla_image_dic[str(nodeid)]={}
        dnsla_image_dic[str(nodeid)]['dname']={'x':x_dname,'y':y_dname}
        dnsla_image_dic[str(nodeid)]['addr']={'x':x_addr,'y':y_addr}
        dnsla_image_dic[str(nodeid)]['name']=Node.objects.get(nb=nodeid).name
        
    #返回需要的数据
    return  HttpResponse(json.dumps(dnsla_image_dic[str(nodeid)]))

def getDnslaOldData(request,reportid):
    global dnsla_image_dic
    report=get_object_or_404(Report, pk=reportid)
    attact_nodes=report.attact_nodes.all()
    ret={}
    for an in attact_nodes:
        if an.nb==0:
            continue
        ret[str(an.nb)]=json.loads(an.dnsla_json)
    #刷新全局变量
    dnsla_image_dic.update(ret)
    return  HttpResponse(json.dumps(ret))

def saveReport(request,report_id):

    report = get_object_or_404(Report, pk=report_id)
    
    image_str=request.POST['image']
    image_base64=image_str.split(',')[1]
    #保存图片
    imgData = base64.b64decode(image_base64)
    leniyimg = open('imgout.png','wb')
    leniyimg.write(imgData)
    leniyimg.close()
    
    from docx import Document
    from docx.shared import Inches,Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    
    document = Document()
    
    create_time=str(report.create_time.date())
    persions=report.persion.split(':::')
    persions=' '.join(persions)
    processes=report.process.split(':::')
    
    
    year=report.start_time.year
    month=report.start_time.month
    day=report.start_time.day
    end_day=report.end_time.day
    
    start_time=str(report.start_time.time())
    end_time=str(report.end_time.time())
    
    attact_nodes=report.attact_nodes.all()
    node_range=''
    for an in attact_nodes:
        try:
            node=Node.objects.get(nb=an.nb)
            an.name=node.name
            node_range+=an.name
            an.abbr=node.abbr
        except:
            continue
    #填写word
    title=document.add_heading(report.title, 0)
    p1=u'填表日期:%s'%create_time
    p = document.add_paragraph(p1)
    #段落格式
    pf = p.paragraph_format
    #向右对齐
    pf.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    #TODO 字体加粗
    run = p.add_run()
    font = run.font
    font.bold=True
    
    
    p = document.add_paragraph(u'参与人员：%s'% persions )
    p = document.add_paragraph(u'故障时间：%d年%d月%d日%s-%d日%s'%(year,month,day,start_time,end_day,end_time))
    p = document.add_paragraph(u'故障类型：域名、网络、安全')
    p = document.add_paragraph(u'影响范围：%s'%node_range)
    p = document.add_heading(u'故障摘要：',level=1)
    p = document.add_paragraph(report.abstract)
    pf = p.paragraph_format
    pf.first_line_indent = Inches(0.25)
    
    p = document.add_heading(u'故障现象：',level=1)
    table = document.add_table(rows=1, cols=3)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = u'节点名称'
    hdr_cells[1].text = u'平均QPS'
    hdr_cells[2].text = u'峰值QPS'
    for an in attact_nodes:
        hdr_cells = table.add_row().cells
        hdr_cells[0].text = an.name
        hdr_cells[1].text = str(an.average_qps)
        hdr_cells[2].text = str(an.max_qps)
    p = document.add_heading(u'处理过程：',level=1)
    for pro in processes:
        if pro == '':
            continue
        document.add_paragraph(
            pro, style='ListBullet'
        )
    p = document.add_heading(u'附录一 MRTG图像',level=1)
    for an in attact_nodes:
        p = document.add_paragraph(an.name)
        document.add_picture(settings.MEDIA_ROOT+an.picture.url, width=Inches(2.5))
    p = document.add_heading(u'附录二 DNSLA图像',level=1)
    document.add_picture('imgout.png', width=Inches(5))
    
    
    
    
    '''
    p = document.add_paragraph('A plain paragraph having some ')
    p.add_run('bold').bold = True
    p.add_run(' and some ')
    p.add_run('italic.').italic = True

    document.add_heading('Heading, level 1', level=1)
    document.add_paragraph('Intense quote', style='IntenseQuote')

    document.add_paragraph(
        'first item in unordered list', style='ListBullet'
    )
    document.add_paragraph(
        'first item in ordered list', style='ListNumber'
    )

    document.add_picture('imgout.png', width=Inches(2.5))
    
    table = document.add_table(rows=1, cols=3)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Qty'
    hdr_cells[1].text = 'Id'
    hdr_cells[2].text = 'Desc'
    for item in recordset:
        row_cells = table.add_row().cells
        row_cells[0].text = str(item.qty)
        row_cells[1].text = str(item.id)
        row_cells[2].text = item.desc
    '''
    document.add_page_break()
    
    filename='%s.docx'%(report.title)
    document.save(settings.MEDIA_ROOT+settings.MEDIA_URL+'tmp-report/'+filename)
    
    
    """
    from django.core.servers.basehttp import FileWrapper
    
    wrapper = FileWrapper(file(filename))
    response = HttpResponse(wrapper, content_type='text/plain')
    response['Content-Length'] = os.path.getsize(filename)
    response['Content-Encoding'] = 'utf-8'
    response['Content-Disposition'] = 'attachment;filename=%s' % filename
    return response
    """
    ret={'filename':filename}
    return HttpResponse(json.dumps(ret))
 
def sendmail(request):
    # if this is a POST request we need to process the form data
    if request.method == 'POST':
        # create a form instance and populate it with data from the request:
        form = ContactForm(request.POST)
        # check whether it's valid:
        if form.is_valid():
            #pdb.set_trace()
            data={"is_success":True,"error_message":'','create_form':False}
            from django.core.mail import send_mail
            subject = form.cleaned_data['subject']
            text = form.cleaned_data['message']
            sender = form.cleaned_data['sender']
            cc_myself = form.cleaned_data['cc_myself']

            #recipients = ['sunyuwu@cnnic.cn']
            #if cc_myself:
                #recipients.append(sender)

            #send_mail(subject, message, sender, recipients)
            try:
                import time,datetime
                import os, smtplib, mimetypes  
                from email.mime.text import MIMEText  
                from email.mime.image import MIMEImage  
                from email.mime.multipart import MIMEMultipart 
                MAIL_LIST = []
                CC=[]
                if cc_myself:
                    CC.append(sender)
                MAIL_FROM = MAIL_USER + "<"+MAIL_USER + "@" + MAIL_POSTFIX + ">" 
                message = MIMEMultipart()  
                message.attach(MIMEText("from %s:\n%s"%(sender,text),'plain','utf-8'))  
                message["Subject"] = subject  
                message["From"] = MAIL_FROM
                message["To"] = ";".join(MAIL_LIST)  
                message["BCC"] = ";".join(BCC)
                message["CC"] = ";".join(CC)
                
                smtp = smtplib.SMTP()  
                smtp.connect(MAIL_HOST)  
                smtp.login(MAIL_USER, MAIL_PASS)  
                smtp.sendmail(MAIL_FROM, MAIL_LIST+BCC+CC, message.as_string())
                smtp.quit() 
            except Exception,e:
                data={"is_success":False,"error_message":str(e),'create_form':False}
            return render(request, 'reports/sendmail.html', data)

    # if a GET (or any other method) we'll create a blank form
    else:
        form = ContactForm()

    return render(request, 'reports/sendmail.html', {'form': form,'create_form':True})

def max_of_list(input_list):
    max_num=0
    for i in xrange(0,len(input_list)):
        if input_list[i]>max_num:
            max_num=input_list[i]
    return max_num

def average_of_list(input_list):
    sum_num=0
    for i in xrange(0,len(input_list)):
        sum_num+=input_list[i]
    if len(input_list)!=0:
        return sum_num/float(len(input_list))
    else:
        return 0

def delete(request,report_id):
    report = get_object_or_404(Report, pk=report_id)
    
    #解除多对多关系
    domains=report.domains.all()
    for obj in domains:
        obj.report.remove(report)
        if len(obj.report.all())==0:
            obj.delete()

    ips=report.ips.all()
    for obj in ips:
        obj.report.remove(report)
        if len(obj.report.all())==0:
            obj.delete()

    #直接删除外键
    attact_nodes=report.attact_nodes.all()
    for obj in attact_nodes:
        #从磁盘上删除文件
        picture_path=os.path.join(settings.MEDIA_ROOT,obj.picture.url[1:])
        if os.path.isfile(picture_path):
            os.remove(picture_path)
        obj.delete()
    
    report.delete()
    
    return index(request)





